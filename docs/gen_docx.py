from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

t = doc.add_heading("社交媒体用户风险智能分析系统 — 技术亮点说明", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("一、系统概述", level=1)
doc.add_paragraph(
    "本模块基于大语言模型（Qwen3.5-27B 本地部署），实现了对社交媒体用户的自动化风险标签判定与用户画像生成，"
    "并支持定时调度、结果回写数据库及钉钉推送通知。"
)

doc.add_heading("二、技术实现亮点", level=1)

doc.add_heading("2.1 两阶段分析架构", level=2)
doc.add_paragraph(
    '系统采用"规则预判 + 大模型归一判定"的两阶段策略，兼顾准确率与效率：'
)
doc.add_paragraph(
    "第一阶段：基于涉政敏感账号名单（23个）和关键词库，快速筛选候选标签，为模型提供分析线索。",
    style="List Bullet",
)
doc.add_paragraph(
    "第二阶段：将用户资料、推文样本、关注关系、候选标签一并送入 Qwen3.5，输出最终标签与用户画像。",
    style="List Bullet",
)
doc.add_paragraph(
    "\u5185\u7f6e\u5173\u952e\u8bcd\u5e93\uff08\u652f\u6301\u6269\u5c55\uff09\uff1a\u6d89\u653f\u7c7b\u5305\u542b\u201c\u4e60\u8fd1\u5e73\u3001\u5206\u88c2\u56fd\u5bb6\u3001\u98a0\u8986\u653f\u6743\u3001\u516d\u56db\u3001\u85cf\u72ec\u3001\u53f0\u72ec\u3001\u6e2f\u72ec\u201d\u7b4912\u4e2a\u8bcd\uff1b"
    "\u6d89\u6050\u7c7b\u5305\u542b\u201c\u6050\u6016\u88ad\u51fb\u3001ISIS\u3001\u5723\u6218\u201d\u7b497\u4e2a\u8bcd\uff1b\u53e6\u542b\u6d89\u8d4c\u3001\u7f51\u7edc\u72af\u7f6a\u3001\u4ec7\u6068\u8a00\u8bba\u4e09\u7c7b\uff0c\u5408\u8ba136\u4e2a\u9ad8\u98ce\u9669\u5173\u952e\u8bcd\u3002"
)

doc.add_heading("2.2 大模型本地私有化部署", level=2)
doc.add_paragraph(
    "Qwen3.5-27B 运行于内网服务器（双卡 RTX 4090），所有用户数据全程不出局域网，保障数据主权。"
    "相比调用公有云 API，本方案无 token 计费限制，支持千级用户批量处理，每用户分析耗时约 70 秒。"
)

doc.add_heading("2.3 URL 短链自动展开与内容增强", level=2)
doc.add_paragraph(
    "推文中大量包含 t.co 短链，系统自动展开短链并抓取目标页面的标题与摘要，拼接到原文后送入模型，"
    "有效补充推文语义信息，提升分类准确率。支持通过代理访问，适配内网环境。"
)

doc.add_heading("2.4 多维度用户画像生成", level=2)
doc.add_paragraph(
    "模型综合账号活跃度、内容风险特征、影响力规模、资料信息等维度，生成 80～220 字的客观用户画像，"
    "风格规范，可直接用于研判报告。画像严格基于真实字段生成，禁止虚构。"
)

doc.add_heading("2.5 定时调度与自动推送", level=2)
doc.add_paragraph(
    "系统支持每天定时（19:00）自动触发增量分析，分析完成后涉政用户自动同步写入 key_focus_user 重点关注表，"
    "并通过钉钉机器人推送任务报告，包含执行统计与涉政用户名单（含主页链接与画像）。"
)

doc.add_heading("三、创新点总结", level=1)
table = doc.add_table(rows=6, cols=2)
table.style = "Table Grid"
table.rows[0].cells[0].text = "创新点"
table.rows[0].cells[1].text = "说明"
rows_data = [
    ("私有化大模型部署", "数据不出内网，安全合规"),
    ("两阶段分析策略", "规则兜底 + 模型语义理解，降低漏判率"),
    ("URL 内容增强", "自动展开短链补充语义，业界较少见"),
    ("增量分析机制", "只处理新增用户，避免重复计算"),
    ("全链路自动化", "从数据读取到结果回写、通知推送，全程无人工干预"),
]
for i, (k, v) in enumerate(rows_data, 1):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.save("docs/系统技术亮点说明.docx")
print("生成成功")
